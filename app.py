import os
import sys
import pandas as pd
import barcode
import tempfile
import base64
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
import streamlit as st
import streamlit.components.v1 as components
import time

# Set up paths
if getattr(sys, 'frozen', False):
    base_path = sys._MEIPASS  # PyInstaller temp folder
else:
    base_path = os.path.dirname(__file__)

downloads_folder = tempfile.gettempdir()
barcode_folder = os.path.join(base_path, "barcodes")

if not os.path.exists(barcode_folder):
    os.makedirs(barcode_folder)

excel_file = os.path.join(base_path, "sku_list.xlsx")

if not os.path.exists(excel_file):
    st.error(f"⚠️ Excel file not found at {excel_file}. Ensure it is present in the project folder.")
    sys.exit(1)

try:
    df = pd.read_excel(excel_file, dtype={"SKU": str})
except Exception as e:
    st.error(f"⚠️ Failed to load Excel file: {e}")
    sys.exit(1)

def generate_barcode(data):
    barcode_path = os.path.join(barcode_folder, data)
    barcode_final_path = barcode_path + ".png"

    if os.path.exists(barcode_final_path):
        return barcode_final_path

    code128 = barcode.get_barcode_class('code128')
    barcode_obj = code128(data, writer=ImageWriter())

    try:
        barcode_obj.save(barcode_path, options={"module_width": 0.25, "module_height": 8})
        if not os.path.exists(barcode_final_path):
            raise FileNotFoundError(f"Barcode file {barcode_final_path} not found after saving!")
        return barcode_final_path
    except Exception as e:
        st.error(f"⚠️ Error generating barcode: {e}")
        sys.exit(1)

def create_label_pdf(sku, description):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(3)
    section.page_height = Inches(1)
    section.top_margin = Inches(0.05)
    section.bottom_margin = Inches(0.05)
    section.left_margin = Inches(0.05)
    section.right_margin = Inches(0.05)

    img_filename = generate_barcode(sku)
    if not os.path.exists(img_filename):
        st.error(f"⚠️ Barcode image {img_filename} missing.")
        sys.exit(1)

    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(img_filename, width=Inches(1.83))
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    desc_para = doc.add_paragraph(description[:34])
    desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    desc_run = desc_para.runs[0]
    desc_run.font.size = Pt(11)
    desc_run.font.name = 'Arial'
    desc_para.paragraph_format.space_before = Pt(0)
    desc_para.paragraph_format.space_after = Pt(0)

    docx_path = os.path.join(downloads_folder, f"{sku}.docx")
    pdf_path = os.path.join(downloads_folder, f"{sku}.pdf")
    doc.save(docx_path)

    try:
        convert(docx_path, pdf_path)
    except Exception as e:
        st.error(f"⚠️ PDF conversion failed: {e}")
        return None

    return pdf_path

# --- Streamlit UI ---
st.title("📦 SKU Barcode Generator")
st.write("Enter an SKU to generate a printable label.")

sku_input = st.text_input("Enter SKU:", "")

if st.button("Generate Label"):
    if sku_input.strip():
        match = df[df["SKU"] == sku_input]
        if match.empty:
            st.warning("⚠️ SKU not found. Please try another.")
        else:
            description = match.iloc[0]["Description"]
            time.sleep(1)
            pdf_path = create_label_pdf(sku_input, description)

            if pdf_path and os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    base64_pdf = base64.b64encode(f.read()).decode("utf-8")

                pdf_viewer = f"""
                <iframe src="data:application/pdf;base64,{base64_pdf}" width="400" height="300"></iframe>
                <br>
                <button onclick="printPDF()">🖨️ Print Label</button>
                <script>
                function printPDF() {{
                    const iframe = document.createElement('iframe');
                    iframe.style.display = 'none';
                    iframe.src = "data:application/pdf;base64,{base64_pdf}";
                    document.body.appendChild(iframe);
                    iframe.onload = function() {{
                        setTimeout(() => {{
                            iframe.contentWindow.focus();
                            iframe.contentWindow.print();
                        }}, 500);
                    }};
                }}
                </script>
                """
                st.success("✅ Label ready for print below.")
                components.html(pdf_viewer, height=400)
            else:
                st.error("❌ Failed to generate label PDF.")
    else:
        st.warning("⚠️ Please enter a valid SKU.")
