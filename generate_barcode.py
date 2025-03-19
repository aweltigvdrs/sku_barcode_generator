import os
import pandas as pd
import barcode
from barcode.writer import ImageWriter
from PIL import ImageFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import shutil
import subprocess
import sys
import tempfile

# Get user's Downloads folder
downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")

# Determine if running as a PyInstaller executable
if getattr(sys, 'frozen', False):
    # Running in a PyInstaller bundle
    base_path = sys._MEIPASS
else:
    # Running as a normal Python script
    base_path = os.path.dirname(__file__)

# Path to the Excel file
excel_file = os.path.join(base_path, "sku_list.xlsx")
print(f"Excel file path: {excel_file}")

# Check if the Excel file exists; exit if not found
if not os.path.exists(excel_file):
    print(f"⚠️ Excel file not found at {excel_file}. Ensure it is bundled or present.")
    sys.exit(1)

# Load SKU list from extracted Excel file
try:
    df = pd.read_excel(excel_file, dtype={"SKU": str})
except Exception as e:
    print(f"⚠️ Failed to load Excel file: {e}")
    sys.exit(1)

# Create a temporary folder for barcodes
barcode_folder = tempfile.mkdtemp()

def generate_barcode(data, filename):
    code128 = barcode.get_barcode_class('code128')
    barcode_obj = code128(data, writer=ImageWriter())
    barcode_path = os.path.splitext(filename)[0]  # Remove extension

    # Explicitly load a TrueType font for text rendering
    font_path = os.path.join(base_path, "arial.ttf")  # Bundled font path (ensure it exists)
    try:
        font = ImageFont.truetype(font_path, size=12)
        barcode_obj.save(barcode_path, options={"module_width": 0.25, "module_height": 8})  # Scaled-down barcode
    except OSError:
        print(f"⚠️ Font file not found or cannot be opened: {font_path}. Ensure it is bundled.")
        sys.exit(1)

def set_margins(section, top=0.05, bottom=0.05, left=0.05, right=0.05):
    section.page_width = Inches(3)
    section.page_height = Inches(1)  # Slightly increased height
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def create_word_doc(sku, description):
    doc = Document()
    section = doc.sections[0]
    set_margins(section)

    img_filename = os.path.join(barcode_folder, f"{sku}.png")
    generate_barcode(sku, img_filename)

    # Add barcode image
    barcode_para = doc.add_paragraph()
    barcode_run = barcode_para.add_run()
    barcode_run.add_picture(img_filename, width=Inches(1.83))  # Scaled-down barcode width
    barcode_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    barcode_para.paragraph_format.space_before = Pt(0)  # Remove space before
    barcode_para.paragraph_format.space_after = Pt(0)  # Remove space after

    # Add description text below barcode
    desc_para = doc.add_paragraph(description[:38])  # Only first 38 characters
    desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    desc_run = desc_para.runs[0]
    desc_run.font.size = Pt(11)
    desc_run.font.name = 'Arial'
    desc_para.paragraph_format.space_before = Pt(0)  # Remove space before
    desc_para.paragraph_format.space_after = Pt(0)  # Remove space after

    # Save the document to Downloads folder
    doc_filename = os.path.join(downloads_folder, f"{sku}.docx")
    doc.save(doc_filename)

    # Open the Word document (Windows-specific)
    subprocess.run(["start", "", doc_filename], shell=True)

    print(f"✅ Label for {sku} generated and saved as {doc_filename}")

# Interactive SKU input loop
while True:
    sku_input = input("Enter SKU (or type 'exit' to quit): ").strip()
    if sku_input.lower() == 'exit':
        break
    
    match = df[df["SKU"] == sku_input]
    if match.empty:
        print("⚠️ SKU not found. Try again.")
        continue

    description = match.iloc[0]["Description"]
    create_word_doc(sku_input, description)

# Cleanup barcode folder once all processing is done
shutil.rmtree(barcode_folder)
